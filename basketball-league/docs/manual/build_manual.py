#!/usr/bin/env python3
"""Builds the Basketball League product manual (.docx) with embedded screenshots."""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
SHOTS = os.path.join(HERE, "screenshots")
ORANGE = RGBColor(0xF3, 0x70, 0x21)

doc = Document()

# Base styles
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ORANGE if level <= 1 else RGBColor(0x33, 0x33, 0x33)
    return h

def para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    return p

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p

def step(n, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p

def shot(filename, caption, width=6.3):
    path = os.path.join(SHOTS, filename)
    if not os.path.exists(path):
        para(f"[missing screenshot: {filename}]", italic=True)
        return
    doc.add_picture(path, width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

# ---------- Cover ----------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Basketball League")
tr.bold = True
tr.font.size = Pt(34)
tr.font.color.rgb = ORANGE
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Product Manual")
sr.font.size = Pt(18)
sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
d = doc.add_paragraph()
d.alignment = WD_ALIGN_PARAGRAPH.CENTER
d.add_run("A guide to running a basketball league: seasons, teams, brackets, live games, and standings.").italic = True
doc.add_paragraph()

# ---------- Intro ----------
heading("About this system", 1)
para("Basketball League runs a season from start to finish. You set up divisions and teams, draw the playoff brackets, schedule games, score them live, stream them, and track the standings. When a season ends it goes into an archive you can look back on.")
para("Three kinds of people use it:")
bullet("Admin. Runs the league. Creates seasons, teams, brackets, schedules, and other accounts.")
bullet("Team manager. Runs one team. Builds the roster and follows the team's schedule. Signs up alone and waits for an admin to approve them.")
bullet("Public viewer. Anyone. Browses teams, schedules, standings, and brackets, and watches live games. No account needed.")

# ---------- Feature list ----------
heading("What the system can do", 1)

heading("Accounts and access", 2)
bullet("Three roles: admin, team manager, and public viewer.")
bullet("Sign in with your email or your username.")
bullet("Admins create admin and manager accounts.")
bullet("Managers sign up on their own, pick or create a team, and wait for an admin to approve them.")

heading("Seasons", 2)
bullet("One season is active at a time.")
bullet("Make a new season as a draft, then activate it when you are ready.")
bullet("Copy divisions, teams, and rosters from an old season. The team's manager moves with it.")
bullet("Set the start date when you activate. Activating the new season ends the old one.")
bullet("End a season to send it to the archive.")
bullet("The archive keeps each finished season for good: divisions, teams, rosters, standings, and brackets. You can read it but not change it.")

heading("Divisions, teams, and players", 2)
bullet("Add as many divisions to a season as you want.")
bullet("Register a team into a division and assign its manager.")
bullet("Give a team a logo. The logo color tints the team in the bracket.")
bullet("Managers fill the roster: name, jersey number, position, height, and photo.")

heading("Brackets", 2)
bullet("Draw a single-elimination bracket on a board.")
bullet("Add match boxes. Two boxes feed the next round.")
bullet("Pick the two teams in each box. Winners move forward on their own.")
bullet("Set a date, time, and venue per match. Remove a match that has not started.")
bullet("Mark one bracket as the division default. New teams join its first round.")
bullet("Save and publish. The public sees a read-only copy.")
bullet("Download any bracket as a picture.")

heading("Schedule and games", 2)
bullet("See every match in the active season. Filter by division.")
bullet("Matches sort by state: live first, then starting, scheduled, and ended.")
bullet("Admins get a separate list of matchups that still need a date.")
bullet("Public viewers and managers see only matches that have a date or are underway. Managers see only their own team.")
bullet("Score a game with the +1, +2, +3, and -1 buttons.")
bullet("A bracket game cannot end in a tie. You enter the final score first.")

heading("Live stream", 2)
bullet("Broadcast a game from a camera or a shared screen.")
bullet("Choose the camera, microphone, and video quality.")
bullet("Viewers watch the stream and see the score change in real time.")
bullet("One person streams a match at a time. Others see who is broadcasting.")

heading("Standings, news, and the public site", 2)
bullet("Win-loss tables per division, built from finished games.")
bullet("Champions and results post to a news feed.")
bullet("Anyone can browse the public site without signing in.")

doc.add_page_break()

# ---------- Getting started ----------
heading("Getting started: signing in", 1)
para("Open the site and sign in with your email or username and password. New team managers click \"Register as a team manager\" to make their own account.")
shot("01-login.png", "The sign-in screen. Managers can register from here.")

# ---------- ADMIN GUIDE ----------
heading("Admin guide", 1)

heading("The admin dashboard", 2)
para("After you sign in, the dashboard shows the count of teams and players in the active season and which season is running. With no active season the counts read zero. The cards link to the rest of the app, and upcoming games and news sit below.")
shot("02-admin-dashboard.png", "The admin dashboard with no active season yet.")

heading("Adding another admin", 2)
para("Open Users. Fill in the name, email, username, and password, set the role to Admin, and click Create.")
shot("03-create-admin.png", "Creating a second admin account.")
para("The new admin shows up in the table below the form.")
shot("04-users-list.png", "The user list after adding an admin.")

heading("Adding a team manager", 2)
para("On the same Users page, set the role to Team Manager and create the account. A manager you make here has no team yet. You assign a team to them when you register the team.")
shot("05-create-manager.png", "Creating a team manager account.")

heading("Creating a season and importing from the past", 2)
para("Open Seasons and click Add a new season. Name it, then pick an old season to copy from. Check the divisions, teams, and rosters you want. Each team can bring its roster with the \"include roster\" box, and \"Select all\" grabs everything.")
shot("06-season-import.png", "Choosing what to copy from a past season.")
para("Click Create. The season starts as a draft with the teams you copied. You can still add divisions by hand.")
shot("07-season-draft-divisions.png", "A new draft season with imported divisions, plus a new one being added.")

heading("Activating a season", 2)
para("A draft does nothing until you activate it. Click Activate season, pick the start date, and confirm. This ends the season that was running before, so only one season is ever live.")
shot("08-activate-dialog.png", "Setting the start date when activating.")
para("The season is now active. From here you manage its divisions, jump to the brackets, or end the season.")
shot("09-season-active.png", "The active season page.")

heading("Looking back at past seasons", 2)
para("Click View season archive to list every finished season.")
shot("10-archive-list.png", "The archive list of ended seasons.")
para("Open one to see its full record: divisions and rosters, the final standings, and the brackets. The archive is read-only, so the history stays put.")
shot("11-archive-detail.png", "A finished season in the archive: rosters, standings, and brackets.")

heading("Adding teams", 2)
para("Open Teams to see every team in the league, grouped by division. Use the dropdown to switch divisions.")
shot("12-teams.png", "Teams, grouped by division.")
para("Click Register Team. Give it a name, pick the division, and choose the manager who will run it. Add a logo if you have one.")
shot("13-register-team.png", "Registering a team and assigning a manager.")
para("The team page shows its manager, roster, and a danger zone for deleting it. Managers fill the roster themselves.")
shot("14-team-detail.png", "A team's page with its manager and an empty roster.")

heading("Building a bracket", 2)
para("Open Brackets. You see the brackets that already exist and a button to make a new one.")
shot("15-brackets-landing.png", "The brackets page.")
para("Make a bracket, pick its division, and name it. On the board, add match boxes and click a slot to drop a team in. Two boxes join into the next round, and the winner moves forward once the game ends. Set a date per match with the clock icon, and remove a match with the trash icon. Click Save & Publish when you are done.")
shot("16-bracket-canvas.png", "The bracket board with teams placed and a final waiting.")
para("The public sees a clean, read-only copy with team logos and colors, and can download it as a picture.")
shot("34-public-bracket.png", "The published bracket the public sees.")

heading("Scheduling matches", 2)
para("Open Schedule. The top table holds games with a date. Below it, admins get a reminder for matchups that have both teams but no date yet.")
shot("17-schedule-admin.png", "The admin schedule, with a reminder about unscheduled games.")
para("Open a match and click Edit schedule to set the date, time, and venue.")
shot("18-schedule-match.png", "Setting a match date and venue.")

heading("Running a game", 2)
para("When a game is underway, the match page shows a scoreboard and a broadcast panel. Tap +1, +2, +3, or -1 to keep score. The score saves on its own.")
shot("19-match-live.png", "The scoreboard and broadcast setup for a live game.")
para("Pick your camera or screen, your microphone, and the quality, then click Go Live. A red LIVE badge and a timer appear while you broadcast.")
shot("20-livestream-host.png", "Broadcasting a game with the score at 5 to 3.")

heading("Standings", 2)
para("Open Standings for the win-loss table of each division, built from finished games.")
shot("22-standings.png", "League standings by division.")

# ---------- MANAGER GUIDE ----------
doc.add_page_break()
heading("Team manager guide", 1)

heading("Signing up", 2)
para("From the sign-in screen, click Register as a team manager. First enter your details. You can have the system generate a strong password.")
shot("23-register-step1.png", "Step one of sign-up: your details.")
para("Next, set up your team. Create a brand-new team in a division, or claim a team that has no manager yet.")
shot("24-register-step2.png", "Step two: create a team or claim an open one.")

heading("Waiting for approval", 2)
para("You are signed in right away, but your account is pending until an admin approves it. The dashboard tells you so.")
shot("25-manager-pending-dashboard.png", "A pending manager's dashboard.")
para("The Players tab shows the same notice until you are approved.")
shot("26-manager-players-pending.png", "The Players tab while waiting for approval.")

heading("Admin approval", 2)
para("An admin opens Users, finds you under Pending registrations, and clicks Approve. That creates or assigns your team and turns your account on.")
shot("27-pending-approval.png", "The admin's pending-registration queue.")

heading("Building your roster", 2)
para("Once approved, open Players and click Add Player. Enter the name, jersey number, position, and height, and add a photo if you like.")
shot("28-add-player.png", "Adding a player to the roster.")
para("Players you add show up on your roster.")
shot("29-manager-roster.png", "A roster after adding a player.")

heading("Your schedule", 2)
para("Your Schedule page shows only your team's games and your division's bracket. There is no division picker, since you only manage one team.")
shot("30-manager-schedule.png", "A manager's schedule, fixed to their own team.")

# ---------- PUBLIC + SETTINGS ----------
doc.add_page_break()
heading("Public viewer guide", 1)
para("Anyone can browse the public site without signing in. The home page shows the latest news and quick links.")
shot("32-public-home.png", "The public home page.")
para("Standings are open to everyone, division by division.")
shot("33-public-standings.png", "Public standings.")
para("On a live match page, viewers watch the stream and see the score change as it happens. If you also help run a team, the page tells you who is broadcasting.")
shot("21-livestream-viewer.png", "A viewer's live match page.")

heading("Settings", 1)
para("Every signed-in user has a Settings page to update their profile and change their password.")
shot("31-settings.png", "The settings page.")

# ---------- Reference ----------
heading("Quick reference: who sees what", 1)
para("The top tabs change with your role.")
bullet("Admin: Dashboard, Teams, Schedule, Standings, Users. Plus Seasons and Brackets from the dashboard and schedule.")
bullet("Team manager: Dashboard, Players, Schedule, Standings.")
bullet("Public viewer: Dashboard, Teams, Schedule, Standings, and Sign in.")

out = os.path.join(HERE, "Basketball-League-Manual.docx")
doc.save(out)
print("Saved", out)
