#!/usr/bin/env python3
"""
Compile the screens/ directory into a single User Manual .docx file.

Run: python3 scripts-build-user-manual.py
Output: screens/USER_MANUAL.docx
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).parent
SCREENS = ROOT / "screens"
OUT = SCREENS / "USER_MANUAL.docx"

# Section, file, title, description
PAGES = [
    # --- Login & onboarding ---
    ("Sign In", "00-login.png", "Sign in",
     "Single login screen for both Admin and Team Manager. Public viewers can browse without an account using the 'Browse without an account' link. Latest announcements are previewed below the login card."),

    # --- Public viewer ---
    ("Public Viewer", "public-01-home.png", "Public home (Mayor's Cup landing page)",
     "Public landing showing season title, summary tiles for total Teams and Matches, an Upcoming Matches list, and a 'View Standings' shortcut. Top nav has Dashboard / Teams / Schedule / Standings links open to public viewers."),
    ("Public Viewer", "public-02-teams.png", "Public teams (grouped by division)",
     "Read-only view of all league teams grouped by division (East, North, South, West). Each card shows the team logo, name, and creation date. Click a card to open the team's roster."),
    ("Public Viewer", "public-03-team-detail.png", "Public team detail",
     "Team hero with logo and division badge, the assigned manager's display name, and the full roster. Each player row shows their jersey number prominently with the position pill on the right. Players are clickable."),
    ("Public Viewer", "public-04-player-detail.png", "Public player detail",
     "Read-only player profile showing jersey number tile (or photo if uploaded), name, position, team and height. No private contact info is exposed publicly."),
    ("Public Viewer", "public-05-schedule.png", "Public schedule + season bracket",
     "Bracket-tree per division with classic box-in-box match cards and connector lines, followed by a full match list table. Match cards link to the public match detail page."),
    ("Public Viewer", "public-06-match-detail.png", "Public match detail",
     "Read-only match page. Status badge auto-updates in real time (polled every 2s). Score board only shows for started/live/ended matches. Live stream player only shows when the match is currently started or live."),
    ("Public Viewer", "public-07-standings.png", "Public standings",
     "Standings tables computed from final results, per division. The 'View Bracketing' banner at top redirects to the schedule page where the full bracket lives."),
    ("Public Viewer", "public-08-announcements.png", "Public announcements feed",
     "Read-only announcements feed showing match results, schedule changes, and championship updates."),

    # --- Admin ---
    ("Admin", "admin-01-dashboard.png", "Admin dashboard",
     "At-a-glance KPIs (Total Teams, Total Players, Active Season). 'View Standings' quick action in the header. Upcoming Matches list with date chips, team-vs-team, and venue. Latest Announcements grid."),
    ("Admin", "admin-02-teams.png", "Teams (admin)",
     "Teams grouped per division with 'Add Division' in the header and 'Register Team' inside each division row. Admins can double-click a division name to rename in place. Each team card shows the team image."),
    ("Admin", "admin-03-team-detail.png", "Team detail (admin)",
     "Manager card (with switch-manager dialog), full roster (jersey number shown beside each name), Edit Team form with team image upload, and a Danger Zone for deletion."),
    ("Admin", "admin-04-schedule.png", "Schedule (admin)",
     "Top: Auto-generate, Create Match, and Open canvas actions. Middle: read-only bracket tree pulled from the canvas. Bottom: full match table with status pills and View links."),
    ("Admin", "admin-05-match-detail-scheduled.png", "Match detail (admin)",
     "Status badge auto-updates in real time. Score board, live stream panel, and schedule editor are conditionally rendered based on the effective status (planned/scheduled/started/live/ended). When live, the stream panel exposes Resume Stream + End Match controls."),
    ("Admin", "admin-06-standings.png", "Standings (admin)",
     "'View Bracketing' banner links to the schedule view. Per-division standings tables computed from finalized matches."),
    ("Admin", "admin-07-users.png", "User management",
     "Create user form for Admin and Team Manager accounts plus a sortable users table with edit/delete actions and team assignment."),
    ("Admin", "admin-08-season-canvas.png", "Season canvas (admin only)",
     "Composes the season bracket per division. Same bracket-tree visual as the public schedule but every match card is clickable to open the MatchPanel for editing scores, status, schedule, advancing winners, marking division/season finals, and elimination overrides. Per-round 'Add to QF/SF/Final' buttons plus 'Add round N' below each tree."),
    ("Admin", "admin-09-announcements.png", "Announcements (admin)",
     "Lists all announcements with a 'New announcement' action. Auto-generated entries (match results, schedule changes, championships) appear alongside any manual posts."),
    ("Admin", "admin-10-settings.png", "Admin settings",
     "Profile edit, Change password (requires current password), Season management section with 'Manage current season' link and 'Start a new season' dialog."),
    ("Admin", "admin-11-user-menu.png", "Top-right user menu (dropdown)",
     "Click the @username chip in the header to open the dropdown: Settings (links to /settings), Log out. Available to both Admin and Manager."),

    # --- Manager ---
    ("Team Manager", "manager-01-dashboard.png", "Manager dashboard",
     "Manager view shows My Roster size, Active Season, Upcoming Matches list, and Latest Announcements. Top nav swaps Teams for Players and hides Users."),
    ("Team Manager", "manager-02-players.png", "Players (manager only)",
     "Manage your team's roster: Add Player, view jersey number / position / height per row, and Remove individual players."),
    ("Team Manager", "manager-03-settings.png", "Manager settings",
     "Profile edit (name, email, username, contact number) and Change password. Manager view does not include the Season management section that admins see."),
]


def add_caption(doc, idx, total, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Figure {idx} of {total}: {title}")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Mayor's Cup Basketball League")
    run.bold = True
    run.font.size = Pt(28)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("User Manual")
    sub_run.font.size = Pt(18)

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro_run = intro.add_run(
        "Web-based league management with public viewer, team-manager portal, and full admin suite. "
        "Includes per-match live streaming via Agora, real-time score board, season bracketing, and "
        "announcement feeds.",
    )
    intro_run.italic = True
    intro_run.font.size = Pt(11)

    doc.add_paragraph()

    # Roles section
    doc.add_heading("Roles", level=1)
    role_table = doc.add_table(rows=4, cols=2)
    role_table.style = "Light Grid Accent 1"
    cells = role_table.rows[0].cells
    cells[0].text = "Role"
    cells[1].text = "Capabilities"
    rows_data = [
        ("Public Viewer",
         "Browse teams, players, schedule, bracket, standings, and announcements without an account."),
        ("Team Manager",
         "Manage own team's roster, view schedule, view standings, edit own profile/password."),
        ("Admin",
         "Full league control: teams, divisions, players, matches, season canvas, users, announcements, broadcaster role for live streams."),
    ]
    for i, (r, cap) in enumerate(rows_data, start=1):
        role_table.rows[i].cells[0].text = r
        role_table.rows[i].cells[1].text = cap

    doc.add_paragraph()

    # Status workflow
    doc.add_heading("Match Status Workflow", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "Matches transition through five statuses. The displayed status is computed at read time so "
        "that 'started' is derived automatically when the current wall-clock hour matches the schedule:",
    )
    bullets = [
        ("Planned", "Match created without a schedule yet."),
        ("Scheduled", "Schedule set, not yet within the start hour."),
        ("Started", "Schedule hour is current; live stream not yet started."),
        ("Live", "Broadcaster is live. Score board is editable by admin."),
        ("Ended", "Admin clicked 'End Match'. Score is final and locked."),
    ]
    for name, desc in bullets:
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(name + ": ")
        run.bold = True
        para.add_run(desc)

    doc.add_paragraph()

    # Default credentials
    doc.add_heading("Default Credentials (dev seed)", level=1)
    p = doc.add_paragraph(
        "Passwords are stored as bcrypt hashes (cost 10). The seed inserts these accounts; "
        "if a user changes their password via /settings, the seed default no longer works."
    )
    cred_table = doc.add_table(rows=6, cols=4)
    cred_table.style = "Light Grid Accent 1"
    headers = ["Role", "Username", "Email", "Password"]
    for i, h in enumerate(headers):
        cred_table.rows[0].cells[i].text = h
    creds = [
        ("admin", "admin", "admin@league.test", "admin123"),
        ("team_manager", "sharks_mgr", "manager@league.test", "manager123"),
        ("team_manager", "warriors_mgr", "warriors@league.test", "manager123"),
        ("team_manager", "eagles_mgr", "eagles@league.test", "manager123"),
        ("team_manager", "bulls_mgr", "bulls@league.test", "manager123"),
    ]
    for i, c in enumerate(creds, start=1):
        for j, v in enumerate(c):
            cred_table.rows[i].cells[j].text = v

    doc.add_page_break()

    # Walkthrough
    doc.add_heading("Walkthrough", level=1)
    last_section = None
    total = len(PAGES)
    for idx, (section, fname, title, desc) in enumerate(PAGES, start=1):
        if section != last_section:
            doc.add_heading(section, level=2)
            last_section = section

        h = doc.add_heading(title, level=3)
        for r in h.runs:
            r.font.size = Pt(13)

        p = doc.add_paragraph(desc)
        for r in p.runs:
            r.font.size = Pt(10.5)

        path = SCREENS / fname
        if not path.exists():
            warn = doc.add_paragraph(f"[missing screenshot: {fname}]")
            warn.runs[0].italic = True
            continue

        pic_par = doc.add_paragraph()
        pic_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pic_par.add_run()
        run.add_picture(str(path), width=Inches(6.6))

        add_caption(doc, idx, total, title)
        doc.add_paragraph()

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
